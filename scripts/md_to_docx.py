from pathlib import Path
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
input_path = ROOT / "StudyNotion_IEEE_Research_Paper.md"
output_path = ROOT / "StudyNotion_IEEE_Research_Paper.docx"

text = input_path.read_text(encoding="utf-8")
doc = Document()

for raw_line in text.splitlines():
    line = raw_line.rstrip()
    if not line:
        doc.add_paragraph("")
        continue

    if line.startswith("### "):
        doc.add_heading(line[4:].strip(), level=3)
    elif line.startswith("## "):
        doc.add_heading(line[3:].strip(), level=2)
    elif line.startswith("# "):
        doc.add_heading(line[2:].strip(), level=1)
    elif line.startswith("- "):
        doc.add_paragraph(line[2:].strip(), style="List Bullet")
    elif line[:2].isdigit() and line[1:3] == ". ":
        doc.add_paragraph(line[3:].strip(), style="List Number")
    else:
        doc.add_paragraph(line)

doc.save(output_path)
print(f"Created: {output_path}")
